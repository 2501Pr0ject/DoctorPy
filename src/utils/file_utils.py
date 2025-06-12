# src/utils/file_utils.py
"""
Utilitaires pour la gestion des fichiers
"""

import os
import json
import csv
import hashlib
import mimetypes
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, BinaryIO, TextIO
from datetime import datetime
import tempfile
import zipfile
import tarfile
import logging

# Imports pour différents formats
import PyPDF2
import pymupdf  # PyMuPDF pour une meilleure extraction PDF
from docx import Document
import yaml

from src.core.config import get_config

logger = logging.getLogger(__name__)

class FileHandler:
    """Gestionnaire de fichiers pour l'assistant pédagogique"""
    
    def __init__(self):
        self.config = get_config()
        self.uploads_dir = self.config.get_uploads_dir()
        self.exports_dir = self.config.get_exports_dir()
        
        # Types MIME supportés
        self.supported_types = {
            'text/plain': ['.txt', '.md', '.py', '.json', '.yaml', '.yml', '.csv'],
            'application/pdf': ['.pdf'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
            'application/json': ['.json'],
            'text/markdown': ['.md'],
            'text/csv': ['.csv'],
            'application/zip': ['.zip'],
            'application/x-tar': ['.tar'],
            'application/gzip': ['.tar.gz', '.gz']
        }
    
    def validate_file(self, file_path: Union[str, Path], max_size: Optional[int] = None) -> Dict[str, Any]:
        """
        Valide un fichier
        
        Args:
            file_path: Chemin vers le fichier
            max_size: Taille maximale en bytes (optionnel)
            
        Returns:
            Dictionnaire avec les informations de validation
        """
        file_path = Path(file_path)
        
        validation_result = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'file_info': {}
        }
        
        # Vérifier l'existence
        if not file_path.exists():
            validation_result['valid'] = False
            validation_result['errors'].append(f"Fichier introuvable: {file_path}")
            return validation_result
        
        # Informations de base
        stat = file_path.stat()
        validation_result['file_info'] = {
            'name': file_path.name,
            'size': stat.st_size,
            'modified': datetime.fromtimestamp(stat.st_mtime),
            'extension': file_path.suffix.lower(),
            'mime_type': mimetypes.guess_type(str(file_path))[0]
        }
        
        # Vérifier la taille
        max_size = max_size or (self.config.security.max_file_size * 1024 * 1024)
        if stat.st_size > max_size:
            validation_result['valid'] = False
            validation_result['errors'].append(
                f"Fichier trop volumineux: {stat.st_size / (1024*1024):.1f}MB "
                f"(max: {max_size / (1024*1024):.1f}MB)"
            )
        
        # Vérifier l'extension
        if not self.config.is_file_allowed(file_path.name):
            validation_result['valid'] = False
            validation_result['errors'].append(
                f"Type de fichier non autorisé: {file_path.suffix}"
            )
        
        # Vérifier le type MIME
        mime_type = validation_result['file_info']['mime_type']
        if mime_type and not self._is_mime_supported(mime_type):
            validation_result['warnings'].append(
                f"Type MIME potentiellement non supporté: {mime_type}"
            )
        
        return validation_result
    
    def _is_mime_supported(self, mime_type: str) -> bool:
        """Vérifie si un type MIME est supporté"""
        return mime_type in self.supported_types
    
    def read_text_file(self, file_path: Union[str, Path], encoding: str = 'utf-8') -> str:
        """
        Lit un fichier texte
        
        Args:
            file_path: Chemin vers le fichier
            encoding: Encodage du fichier
            
        Returns:
            Contenu du fichier
        """
        file_path = Path(file_path)
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            # Essayer avec d'autres encodages
            for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        content = f.read()
                        logger.warning(f"Fichier lu avec l'encodage {enc} au lieu de {encoding}")
                        return content
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Impossible de décoder le fichier {file_path}")
    
    def read_pdf_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Lit un fichier PDF
        
        Args:
            file_path: Chemin vers le fichier PDF
            
        Returns:
            Dictionnaire avec le contenu et les métadonnées
        """
        file_path = Path(file_path)
        
        result = {
            'text': '',
            'metadata': {},
            'pages': [],
            'images': []
        }
        
        try:
            # Utiliser PyMuPDF pour une meilleure extraction
            doc = pymupdf.open(str(file_path))
            
            result['metadata'] = {
                'page_count': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', '')
            }
            
            full_text = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extraire le texte
                page_text = page.get_text()
                full_text.append(page_text)
                
                result['pages'].append({
                    'page_number': page_num + 1,
                    'text': page_text,
                    'text_length': len(page_text)
                })
                
                # Extraire les images (optionnel)
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    result['images'].append({
                        'page_number': page_num + 1,
                        'image_index': img_index,
                        'width': img[2],
                        'height': img[3]
                    })
            
            result['text'] = '\n\n'.join(full_text)
            doc.close()
            
        except Exception as e:
            logger.error(f"Erreur lors de la lecture du PDF {file_path}: {e}")
            # Fallback avec PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    
                    result['metadata']['page_count'] = len(reader.pages)
                    
                    full_text = []
                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        full_text.append(page_text)
                        
                        result['pages'].append({
                            'page_number': page_num + 1,
                            'text': page_text,
                            'text_length': len(page_text)
                        })
                    
                    result['text'] = '\n\n'.join(full_text)
                    
            except Exception as fallback_error:
                logger.error(f"Erreur fallback PyPDF2: {fallback_error}")
                raise ValueError(f"Impossible de lire le PDF: {e}")
        
        return result
    
    def read_docx_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Lit un fichier DOCX
        
        Args:
            file_path: Chemin vers le fichier DOCX
            
        Returns:
            Dictionnaire avec le contenu et les métadonnées
        """
        file_path = Path(file_path)
        
        try:
            doc = Document(str(file_path))
            
            # Extraire le texte des paragraphes
            paragraphs = []
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text.append(para.text)
            
            # Extraire le texte des tableaux
            tables_content = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append('\t'.join(row_text))
                tables_content.append('\n'.join(table_text))
                full_text.extend(table_text)
            
            result = {
                'text': '\n\n'.join(full_text),
                'paragraphs': paragraphs,
                'tables': tables_content,
                'metadata': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables_content)
                }
            }
            
            # Métadonnées du document si disponibles
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                result['metadata'].update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': props.created,
                    'modified': props.modified
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Erreur lors de la lecture du DOCX {file_path}: {e}")
            raise ValueError(f"Impossible de lire le fichier DOCX: {e}")
    
    def read_json_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Lit un fichier JSON
        
        Args:
            file_path: Chemin vers le fichier JSON
            
        Returns:
            Contenu JSON
        """
        file_path = Path(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except json.JSONDecodeError as e:
            logger.error(f"Erreur JSON dans {file_path}: {e}")
            raise ValueError(f"Fichier JSON invalide: {e}")
    
    def read_csv_file(self, file_path: Union[str, Path], delimiter: str = ',') -> Dict[str, Any]:
        """
        Lit un fichier CSV
        
        Args:
            file_path: Chemin vers le fichier CSV
            delimiter: Délimiteur utilisé
            
        Returns:
            Dictionnaire avec les données et métadonnées
        """
        file_path = Path(file_path)
        
        try:
            # Détecter automatiquement le délimiteur
            with open(file_path, 'r', encoding='utf-8') as f:
                sample = f.read(1024)
                sniffer = csv.Sniffer()
                try:
                    detected_delimiter = sniffer.sniff(sample).delimiter
                    delimiter = detected_delimiter
                except:
                    pass  # Garder le délimiteur par défaut
            
            # Lire le CSV
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                rows = list(reader)
                
                result = {
                    'data': rows,
                    'headers': reader.fieldnames,
                    'metadata': {
                        'row_count': len(rows),
                        'column_count': len(reader.fieldnames) if reader.fieldnames else 0,
                        'delimiter': delimiter
                    }
                }
                
                return result
                
        except Exception as e:
            logger.error(f"Erreur lors de la lecture du CSV {file_path}: {e}")
            raise ValueError(f"Impossible de lire le fichier CSV: {e}")
    
    def read_yaml_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Lit un fichier YAML
        
        Args:
            file_path: Chemin vers le fichier YAML
            
        Returns:
            Contenu YAML
        """
        file_path = Path(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        except yaml.YAMLError as e:
            logger.error(f"Erreur YAML dans {file_path}: {e}")
            raise ValueError(f"Fichier YAML invalide: {e}")
    
    def process_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Traite un fichier selon son type
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Contenu traité avec métadonnées
        """
        file_path = Path(file_path)
        
        # Valider le fichier
        validation = self.validate_file(file_path)
        if not validation['valid']:
            raise ValueError(f"Fichier invalide: {', '.join(validation['errors'])}")
        
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                content = self.read_pdf_file(file_path)
            elif extension == '.docx':
                content = self.read_docx_file(file_path)
            elif extension == '.json':
                json_data = self.read_json_file(file_path)
                content = {
                    'text': json.dumps(json_data, indent=2, ensure_ascii=False),
                    'structured_data': json_data,
                    'metadata': {'type': 'json'}
                }
            elif extension == '.csv':
                csv_data = self.read_csv_file(file_path)
                content = {
                    'text': self._csv_to_text(csv_data),
                    'structured_data': csv_data,
                    'metadata': csv_data['metadata']
                }
            elif extension in ['.yaml', '.yml']:
                yaml_data = self.read_yaml_file(file_path)
                content = {
                    'text': yaml.dump(yaml_data, default_flow_style=False, allow_unicode=True),
                    'structured_data': yaml_data,
                    'metadata': {'type': 'yaml'}
                }
            else:
                # Fichier texte simple
                text_content = self.read_text_file(file_path)
                content = {
                    'text': text_content,
                    'metadata': {'type': 'text', 'encoding': 'utf-8'}
                }
            
            # Ajouter les informations générales du fichier
            content['file_info'] = validation['file_info']
            content['file_path'] = str(file_path)
            content['processed_at'] = datetime.now().isoformat()
            
            return content
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement de {file_path}: {e}")
            raise
    
    def _csv_to_text(self, csv_data: Dict[str, Any]) -> str:
        """Convertit des données CSV en texte lisible"""
        if not csv_data['data']:
            return "Fichier CSV vide"
        
        text_parts = []
        text_parts.append(f"Fichier CSV avec {csv_data['metadata']['row_count']} lignes et {csv_data['metadata']['column_count']} colonnes")
        text_parts.append(f"Colonnes: {', '.join(csv_data['headers'])}")
        text_parts.append("")
        
        # Afficher quelques lignes d'exemple
        for i, row in enumerate(csv_data['data'][:5]):
            row_text = []
            for header in csv_data['headers']:
                value = row.get(header, '')
                row_text.append(f"{header}: {value}")
            text_parts.append(f"Ligne {i+1}: {', '.join(row_text)}")
        
        if len(csv_data['data']) > 5:
            text_parts.append(f"... et {len(csv_data['data']) - 5} autres lignes")
        
        return "\n".join(text_parts)
    
    def save_file(self, content: Union[str, bytes], filename: str, 
                  user_id: str = "default", subfolder: str = "") -> Path:
        """
        Sauvegarde un fichier
        
        Args:
            content: Contenu à sauvegarder
            filename: Nom du fichier
            user_id: ID de l'utilisateur
            subfolder: Sous-dossier optionnel
            
        Returns:
            Chemin du fichier sauvegardé
        """
        # Créer le répertoire utilisateur
        user_dir = self.uploads_dir / user_id
        if subfolder:
            user_dir = user_dir / subfolder
        user_dir.mkdir(parents=True, exist_ok=True)
        
        # Chemin complet du fichier
        file_path = user_dir / filename
        
        # Sauvegarder selon le type de contenu
        if isinstance(content, bytes):
            with open(file_path, 'wb') as f:
                f.write(content)
        else:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        logger.info(f"Fichier sauvegardé: {file_path}")
        return file_path
    
    def export_content(self, content: str, filename: str, format_type: str = "txt",
                      user_id: str = "default") -> Path:
        """
        Exporte du contenu dans différents formats
        
        Args:
            content: Contenu à exporter
            filename: Nom de base du fichier
            format_type: Type d'export (txt, md, json, pdf)
            user_id: ID de l'utilisateur
            
        Returns:
            Chemin du fichier exporté
        """
        # Créer le répertoire d'export
        export_dir = self.exports_dir / user_id
        export_dir.mkdir(parents=True, exist_ok=True)
        
        # Ajouter l'extension selon le format
        if not filename.endswith(f'.{format_type}'):
            filename = f"{filename}.{format_type}"
        
        file_path = export_dir / filename
        
        if format_type == "txt":
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        elif format_type == "md":
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        elif format_type == "json":
            # Essayer de parser le contenu comme JSON, sinon l'encapsuler
            try:
                json_content = json.loads(content)
            except:
                json_content = {"content": content, "exported_at": datetime.now().isoformat()}
            
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(json_content, f, indent=2, ensure_ascii=False)
        
        elif format_type == "pdf":
            try:
                from fpdf import FPDF
                
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                
                # Découper le contenu en lignes
                lines = content.split('\n')
                for line in lines:
                    # Encoder pour éviter les erreurs Unicode
                    try:
                        pdf.cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                    except:
                        pdf.cell(0, 10, "Caractère non supporté", ln=True)
                
                pdf.output(str(file_path))
                
            except ImportError:
                logger.warning("fpdf2 non installé, export en texte à la place")
                file_path = file_path.with_suffix('.txt')
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
        
        logger.info(f"Contenu exporté: {file_path}")
        return file_path
    
    def get_file_hash(self, file_path: Union[str, Path]) -> str:
        """
        Calcule le hash MD5 d'un fichier
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Hash MD5 du fichier
        """
        hash_md5 = hashlib.md5()
        
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        
        return hash_md5.hexdigest()
    
    def create_backup(self, source_path: Union[str, Path], 
                     backup_dir: Optional[Path] = None) -> Path:
        """
        Crée une sauvegarde d'un fichier ou dossier
        
        Args:
            source_path: Chemin source
            backup_dir: Répertoire de sauvegarde (optionnel)
            
        Returns:
            Chemin de la sauvegarde
        """
        source_path = Path(source_path)
        
        if backup_dir is None:
            backup_dir = self.config.base_dir / "backups"
        
        backup_dir.mkdir(parents=True, exist_ok=True)
        
        # Nom de la sauvegarde avec timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"{source_path.stem}_{timestamp}"
        
        if source_path.is_file():
            backup_path = backup_dir / f"{backup_name}{source_path.suffix}"
            shutil.copy2(source_path, backup_path)
        else:
            backup_path = backup_dir / f"{backup_name}.zip"
            with zipfile.ZipFile(backup_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in source_path.rglob('*'):
                    if file_path.is_file():
                        zipf.write(file_path, file_path.relative_to(source_path))
        
        logger.info(f"Sauvegarde créée: {backup_path}")
        return backup_path
    
    def clean_old_files(self, directory: Union[str, Path], 
                       max_age_days: int = 30, pattern: str = "*") -> int:
        """
        Nettoie les anciens fichiers d'un répertoire
        
        Args:
            directory: Répertoire à nettoyer
            max_age_days: Âge maximum en jours
            pattern: Pattern des fichiers à considérer
            
        Returns:
            Nombre de fichiers supprimés
        """
        directory = Path(directory)
        
        if not directory.exists():
            return 0
        
        cutoff_time = datetime.now().timestamp() - (max_age_days * 24 * 3600)
        deleted_count = 0
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                try:
                    file_path.unlink()
                    deleted_count += 1
                    logger.info(f"Fichier supprimé: {file_path}")
                except Exception as e:
                    logger.error(f"Erreur lors de la suppression de {file_path}: {e}")
        
        return deleted_count
    
    def extract_archive(self, archive_path: Union[str, Path], 
                       extract_to: Optional[Path] = None) -> Path:
        """
        Extrait une archive
        
        Args:
            archive_path: Chemin vers l'archive
            extract_to: Répertoire d'extraction (optionnel)
            
        Returns:
            Répertoire d'extraction
        """
        archive_path = Path(archive_path)
        
        if extract_to is None:
            extract_to = archive_path.parent / archive_path.stem
        
        extract_to.mkdir(parents=True, exist_ok=True)
        
        if archive_path.suffix.lower() == '.zip':
            with zipfile.ZipFile(archive_path, 'r') as zipf:
                zipf.extractall(extract_to)
        
        elif archive_path.suffix.lower() in ['.tar', '.tar.gz', '.tgz']:
            with tarfile.open(archive_path, 'r:*') as tarf:
                tarf.extractall(extract_to)
        
        else:
            raise ValueError(f"Format d'archive non supporté: {archive_path.suffix}")
        
        logger.info(f"Archive extraite dans: {extract_to}")
        return extract_to
    
    def get_directory_size(self, directory: Union[str, Path]) -> int:
        """
        Calcule la taille d'un répertoire
        
        Args:
            directory: Répertoire à analyser
            
        Returns:
            Taille en bytes
        """
        directory = Path(directory)
        total_size = 0
        
        for file_path in directory.rglob('*'):
            if file_path.is_file():
                total_size += file_path.stat().st_size
        
        return total_size
    
    def list_files(self, directory: Union[str, Path], 
                  recursive: bool = True, include_hidden: bool = False) -> List[Dict[str, Any]]:
        """
        Liste les fichiers d'un répertoire
        
        Args:
            directory: Répertoire à analyser
            recursive: Recherche récursive
            include_hidden: Inclure les fichiers cachés
            
        Returns:
            Liste des fichiers avec leurs métadonnées
        """
        directory = Path(directory)
        files = []
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory.glob(pattern):
            if file_path.is_file():
                # Ignorer les fichiers cachés si demandé
                if not include_hidden and file_path.name.startswith('.'):
                    continue
                
                stat = file_path.stat()
                
                files.append({
                    'name': file_path.name,
                    'path': str(file_path),
                    'size': stat.st_size,
                    'modified': datetime.fromtimestamp(stat.st_mtime),
                    'extension': file_path.suffix.lower(),
                    'mime_type': mimetypes.guess_type(str(file_path))[0],
                    'relative_path': str(file_path.relative_to(directory))
                })
        
        return sorted(files, key=lambda x: x['name'])


def safe_filename(filename: str) -> str:
    """
    Crée un nom de fichier sûr
    
    Args:
        filename: Nom de fichier original
        
    Returns:
        Nom de fichier sécurisé
    """
    # Caractères interdits dans les noms de fichiers
    forbidden_chars = '<>:"/\\|?*'
    
    # Remplacer les caractères interdits
    safe_name = filename
    for char in forbidden_chars:
        safe_name = safe_name.replace(char, '_')
    
    # Limiter la longueur
    if len(safe_name) > 200:
        name, ext = os.path.splitext(safe_name)
        safe_name = name[:200-len(ext)] + ext
    
    return safe_name


def get_file_encoding(file_path: Union[str, Path]) -> str:
    """
    Détecte l'encodage d'un fichier texte
    
    Args:
        file_path: Chemin vers le fichier
        
    Returns:
        Encodage détecté
    """
    try:
        import chardet
        
        with open(file_path, 'rb') as f:
            raw_data = f.read(10000)  # Lire les premiers 10KB
            result = chardet.detect(raw_data)
            return result['encoding'] or 'utf-8'
    
    except ImportError:
        # Fallback sans chardet
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    f.read(1000)  # Essayer de lire un peu
                return encoding
            except UnicodeDecodeError:
                continue
        
        return 'utf-8'  # Défaut


def format_file_size(size_bytes: int) -> str:
    """
    Formate une taille de fichier en unités lisibles
    
    Args:
        size_bytes: Taille en bytes
        
    Returns:
        Taille formatée
    """
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    
    return f"{size_bytes:.1f} {size_names[i]}"


def create_temp_file(content: Union[str, bytes], suffix: str = ".tmp") -> Path:
    """
    Crée un fichier temporaire
    
    Args:
        content: Contenu du fichier
        suffix: Extension du fichier
        
    Returns:
        Chemin du fichier temporaire
    """
    with tempfile.NamedTemporaryFile(mode='wb' if isinstance(content, bytes) else 'w',
                                   suffix=suffix, delete=False) as tmp_file:
        if isinstance(content, bytes):
            tmp_file.write(content)
        else:
            tmp_file.write(content)
        
        return Path(tmp_file.name)